from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from markdown import markdown
from bs4 import BeautifulSoup
from typing import Any, List, Tuple
import requests
from docifyai.core import logger
import random
import string
import tempfile
from html2docx import html2docx

logger = logger.Logger(__name__)


class Aidoc:
    def __init__(self,
                 repo_data: Any,
                 temp_dir: str,
                 doc_contents: List[Tuple[str, str]],
                 intro: str) -> None:
        self.repo_data = repo_data
        self.doc_contents = doc_contents
        self.project_intro = intro

        # we may use non default template in future
        self.doc = Document()

    @staticmethod
    def gen_unique_name(prefix: str, length: int = 10) -> str:
        random_part = ''.join(random.choices(string.ascii_letters + string.digits, k=length))
        return f"{prefix}{random_part}.docx"

    def create_document(self) -> str:
        self.add_first_page()
        self.add_content_page()
        self.add_contents()
        prefix = self.repo_data["name"]
        doc_temp_file = tempfile.mktemp(".docx", prefix)
        self.doc.save(doc_temp_file)
        return doc_temp_file

    def add_first_page(self) -> None:
        doc = self.doc

        title_head = doc.add_heading(self.repo_data["name"], 0)
        title_head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_para = doc.add_paragraph(self.repo_data["description"])
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # for project info, need to change for proper format
        # should use github key
        """for future we can build a separate callable for calling github api endpoint"""
        lang_res = requests.get(self.repo_data["languages_url"])
        lang_used = {}
        lang_str = ""
        if lang_res.status_code == 200:
            lang_used = lang_res.json()
            for lang, count in lang_used.items():
                lang_str = f"{lang_str}{lang}:{count}\n"

        """Improve the formatting of the project info
        also add avatar and licence type"""

        owner_name = {self.repo_data["owner"]["login"]}
        fork_count = {self.repo_data["forks_count"]}
        info_for_doc = f"\n\n\n\n\n\n\n\n\n\nOwner:{owner_name}\nLanguages Used:\t{lang_str}\nFork Count: {fork_count}"
        footer_para = doc.add_paragraph(info_for_doc)
        footer_para.runs[0].font.size = Pt(12)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_content_page(self) -> None:
        self.doc.add_page_break()
        title_head = self.doc.add_heading("Contents", 0)
        title_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for name, _ in self.doc_contents:
            self.doc.add_heading(name, level=4)
            # add one line of info about the chapter

    def add_contents(self) -> None:
        # add intro page
        doc = self.doc

        # add contents of each chapters
        for name, contents in self.doc_contents:
            doc.add_page_break()
            chapter_head = doc.add_heading(name)
            chapter_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            html_content = markdown(contents)

            # testing html rendering
            # same object is being passed so no problem if we catch doc from html2docx
            doc = html2docx(html_content, doc)
            # doc.add_paragraph(text_content)

            # soup = BeautifulSoup(html_content, "html.parser")

            # processed_elements = set()

            # for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']):
            #     if element.text and element.text not in processed_elements:  # Check for text and avoid duplicates
            #         # Handle headings
            #         if element.name.startswith('h'):
            #             style = f'Heading {element.name[1:]}'
            #             doc.add_paragraph(element.text, style=style)
            #             processed_elements.add(element.text)  # Add processed heading text
            #
            #         # Handle lists (extract only direct text content of list items)
            #         elif element.name in ['ul', 'ol']:
            #             list_element = doc.add_paragraph()
            #             list_element.style = 'List Bullet' if element.name == 'ul' else 'List Number'
            #             for item in element.find_all('li'):
            #                 if item.text and item.text not in processed_elements:  # Check for text and avoid duplicates in list items
            #                     doc.add_paragraph(item.text.strip())  # Extract and add trimmed text
            #                     processed_elements.add(item.text)  # Add processed list item text
            #
            #         # Handle paragraphs (avoid duplicates)
            #         else:
            #             if element.text and element.text not in processed_elements:
            #                 doc.add_paragraph(element.text)
            #                 processed_elements.add(element.text)  # Add processed paragraph text
