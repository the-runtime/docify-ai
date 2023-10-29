from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
from pathlib import Path
from docifyai.core import logger
from docifyai.utils import utils

logger = logger.Logger(__name__)


class ContentGen:
    def __init__(self, doc: Any, root_directory: str, max_depth: int = 2):
        self.doc = doc
        self.root_directory = Path(root_directory)
        self.max_depth = max_depth

    def add_content_page(self):
        # start at new page in doc
        # doc.counter_to_new_page(
        title_head = self.doc.add_heading("Contents", 0)
        title_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # title_run = title_head.runs[0]
        # title_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._generate_contents(self.root_directory)

    def _generate_contents(self, directory: Path, depth: int = 0) -> None:
        if depth > self.max_depth:
            return
        if utils.should_ignore(directory):
            return

        if directory.is_dir():
            children = sorted(
                [
                    child
                    for child in directory.iterdir()
                    if child.name != ".git"
                ]
            )

            for index, child in enumerate(children):
                if utils.should_ignore(child):
                    continue
                if depth == 0:
                    self.doc.add_heading(str(child), level=4)
                elif depth == 1:
                    self.doc.add_paragraph(str(child).rstrip("/").split("/")[-1])
                is_last_child = index == len(children) - 1
                self._generate_contents(child, depth + 1)
