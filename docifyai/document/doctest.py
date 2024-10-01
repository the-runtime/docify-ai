import docx
from bs4 import BeautifulSoup

def html_to_docx(html_content, docx_file):
    """Converts HTML content to a DOCX document.

    Args:
        html_content (str): The HTML content to convert.
        docx_file (str): The filename for the output DOCX document.
    """

    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new DOCX document
    doc = docx.Document()

    # Iterate over each block-level element in the HTML
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']):
        # Handle headings
        if element.name.startswith('h'):
            style = f'Heading {element.name[1:]}'
            doc.add_paragraph(element.text, style=style)
        # Handle lists
        elif element.name in ['ul', 'ol']:
            list_element = doc.add_paragraph()
            list_element.style = 'List Bullet' if element.name == 'ul' else 'List Number'
            for item in element.find_all('li'):
                for run in item.find_all(text=True):  # Find all text runs within the list item
                    run.style = 'List Bullet'  # Apply bullet or numbering style to each run
        # Handle paragraphs
        else:
            doc.add_paragraph(element.text)

    # Save the DOCX document
    doc.save(docx_file)
if __name__ == '__main__':
        # Example usage
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
        <title>Example HTML</title>
        </head>
        <body>
        <h1>This is a heading</h1>
        <p>This is a paragraph.</p>
        <ul>
        <li>Item 1</li>
        <li>Item 2</li>
        </ul>
        </body>
        </html>
        """

        docx_file = "output.docx"
        html_to_docx(html_content, docx_file)