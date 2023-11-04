from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from typing import Dict, Any, Tuple
import requests
from pathlib import Path
from docifyai.core import logger
from docifyai.document import contents

logger = logger.Logger(__name__)


class Aidoc:
    def __init__(self,
                 repo_data: Any,
                 temp_dir: str,
                 code_details: Dict[str, str],
                 folder_details: Dict[str, str],
                 intro_tuple: Tuple[str, str]
                 ) -> None:
        # self.code_summary_data = repo_summary
        self.temp_dir = temp_dir
        self.repo_data = repo_data
        self.intro_content = intro_tuple[1]
        self.code_details = code_details
        self.chapter_details = folder_details

        # we may use non default template in future
        self.doc = Document()

    def create_document(self) -> str:
        self.add_first_page()
        self.add_content_page()
        self.add_intro_page()
        self.add_chapters()
        self.doc.save("documentFileFirst.docx")
        return "documentFileFirst.docx"

    def add_first_page(self) -> None:
        """"""
        doc = self.doc
        repo_data = self.repo_data
        title_head = doc.add_heading(repo_data["name"], 0)
        title_head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_para = doc.add_paragraph(repo_data["description"])
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # for project info
        lang_res = requests.get(repo_data["languages_url"])
        lang_used = {}
        if lang_res.status_code == 200:
            lang_used = lang_res.json()
        lang_str = ""
        for lang, count in lang_used.items():
            lang_str = f"{lang_str}{lang}:{count}\n"

        """Improve the formatting of the project info"""
        """also add avatar and licence type"""

        # section.footer_distance = Pt(36)
        # section.footer_line = False

        # footer = section.footer
        owner_name = {repo_data["owner"]["login"]}
        fork_count = {repo_data["forks_count"]}
        info_for_doc = f"\n\n\n\n\n\n\n\n\n\n\nOwner:{owner_name}\nLanguages Used:\t{lang_str}\nFork Count: {fork_count}"
        footer_para = doc.add_paragraph(info_for_doc)
        footer_para.runs[0].font.size = Pt(12)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # doc.save("documentFileFirst.docx")

        # return "documentFileFirst.docx"

    def add_content_page(self) -> None:
        self.doc.add_page_break()
        title_head = self.doc.add_heading("Contents", 0)
        title_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for name, _ in self.chapter_details.items():
            self.doc.add_heading(name.split('/')[-1], level=4)

    #         add one line of info about the chapter

    def add_intro_page(self) -> None:
        self.doc.add_page_break()
        titlehead = self.doc.add_heading("Introduction", 0)
        titlehead.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(self.intro_content)

    def add_chapters(self) -> None:
        for name, chapter_intro in self.chapter_details.items():
            self.doc.add_page_break()
            chapter_head = self.doc.add_heading(name.split("/")[-1], 0)
            chapter_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph(chapter_intro)
            self.doc.add_page_break()
            if Path(self.temp_dir).joinpath(name).is_file():
                continue
            else:
                actual_path = Path(self.temp_dir).joinpath(name)
                self.recursive_cont_addition(actual_path, 2)
                # logger.debug("folder found")

    def recursive_cont_addition(self, path: Path, level) -> None:
        if path.is_dir():
            for child in path.iterdir():
                if child.is_file():
                    topic_content = self.code_details.get(str(child.relative_to(self.temp_dir)))
                    if topic_content:
                        self.doc.add_heading(str(child).split("/")[-1], level=level)
                        self.doc.add_paragraph(topic_content)
                else:
                    self.recursive_cont_addition(child, level=level + 1)
